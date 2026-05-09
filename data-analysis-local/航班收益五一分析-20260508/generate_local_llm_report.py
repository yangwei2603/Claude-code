from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

# ─── 数据准备 ───────────────────────────────────────────
file = '/Users/fox/Claude Code/data-analysis-local/航班收益五一分析-20260508/航班收益快报明细.xlsx'
df = pd.read_excel(file, sheet_name='航班收益快报明细')

may_df = df[(df['航班日期(YYYY-MM-DD)'] >= '2026-04-30') &
            (df['航班日期(YYYY-MM-DD)'] <= '2026-05-06')]
may_last = df[(df['航班日期(YYYY-MM-DD)'] >= '2025-04-30') &
              (df['航班日期(YYYY-MM-DD)'] <= '2025-05-06')]

total_rev = may_df['总收入'].sum()
total_profit = may_df['航班收益'].sum()
trans_profit = may_df['运输收益'].sum()
total_pax = int(may_df['总人数'].sum())
total_flights = int(may_df['架次'].sum())
avg_load = may_df['客座率'].mean()
days = 7

last_rev = may_last['总收入'].sum()
last_pax = int(may_last['总人数'].sum())
last_flights = int(may_last['架次'].sum())
last_profit = may_last['航班收益'].sum()
last_load = may_last['客座率'].mean()

rev_chg    = (total_rev - last_rev) / last_rev * 100
pax_chg    = (total_pax - last_pax) / last_pax * 100
flight_chg = (total_flights - last_flights) / last_flights * 100
profit_chg = (total_profit - last_profit) / abs(last_profit) * 100

route_nat = may_df.groupby('航线性质：1-国内，2- 区域， 3-国际').agg(
    收入=('总收入','sum'), 架次=('架次','sum'),
    人数=('总人数','sum'), 客座率=('客座率','mean')).reset_index()
route_nat['性质'] = route_nat['航线性质：1-国内，2- 区域， 3-国际'].map({1:'国内',2:'区域',3:'国际'})

daily = may_df.groupby('航班日期(YYYY-MM-DD)').agg(
    架次=('架次','sum'), 收入=('总收入','sum'),
    人数=('总人数','sum'), 客座率=('客座率','mean'),
    收益=('航班收益','sum')).reset_index().sort_values('航班日期(YYYY-MM-DD)')

top10 = may_df.groupby('航线名称').agg(
    收入=('总收入','sum'), 架次=('架次','sum'),
    人数=('总人数','sum'), 客座率=('客座率','mean'),
    航班收益=('航班收益','sum')).sort_values('收入', ascending=False).head(10).reset_index()

base_top5 = may_df.groupby('基地名称').agg(
    收入=('总收入','sum'), 架次=('架次','sum'),
    人数=('总人数','sum'), 客座率=('客座率','mean')).sort_values('收入', ascending=False).head(5).reset_index()

intl = may_df[may_df['航线性质：1-国内，2- 区域， 3-国际'].isin([2,3])]
intl_top = intl.groupby('航线名称').agg(
    收入=('总收入','sum'), 架次=('架次','sum'),
    人数=('总人数','sum'), 客座率=('客座率','mean'),
    航班收益=('航班收益','sum')).sort_values('收入', ascending=False).head(10).reset_index()

var_cost = may_df['变动成本'].sum()
fix_cost = may_df['固定成本费用'].sum()
fuel = may_df['航油费'].sum()
landing = may_df['起降费=航路费＋机场基本收费＋旅客旅客服务费＋旅客安检费＋货邮安检费'].sum()
maint = may_df['维修成本金额= 轮档时间*维修成本标准'].sum()
total_cost = var_cost + fix_cost

depart    = df[(df['航班日期(YYYY-MM-DD)'] >= '2026-04-30') & (df['航班日期(YYYY-MM-DD)'] <= '2026-05-01')]
return_df = df[(df['航班日期(YYYY-MM-DD)'] >= '2026-05-05') & (df['航班日期(YYYY-MM-DD)'] <= '2026-05-06')]
mid_df    = may_df[(may_df['航班日期(YYYY-MM-DD)'] >= '2026-05-02') & (may_df['航班日期(YYYY-MM-DD)'] <= '2026-05-03')]

def fmt(n):
    if abs(n) >= 1e8:   return f"{n/1e8:.1f}亿"
    if abs(n) >= 1e4:   return f"{n/1e4:.0f}万"
    return f"{n:,.0f}"

RED    = RGBColor(0xC0, 0x39, 0x2B)
BLUE   = RGBColor(0x1A, 0x52, 0x76)
LBLUE  = RGBColor(0x29, 0x80, 0xB9)
LGRAY  = RGBColor(0xF2, 0xF3, 0xF4)
MGRAY  = RGBColor(0x7F, 0x8C, 0x8D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREEN  = RGBColor(0x1E, 0x84, 0x49)
DKGRAY = RGBColor(0x2C, 0x3E, 0x50)

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A5276')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = LBLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def note(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size  = Pt(8)
        run.font.italic = True
        run.font.color.rgb = MGRAY
    return p

def spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return p

def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    if col_widths is None:
        col_widths = [Cm(3.5)] * n_cols
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        c.width = col_widths[i]
        set_cell_bg(c, BLUE)
        cell_text(c, h, bold=True, size=9, color=WHITE)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = RGBColor(0xF8, 0xF9, 0xFA) if ri % 2 == 1 else WHITE
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.width = col_widths[ci]
            set_cell_bg(c, bg)
            cell_text(c, val, size=9)
    return table

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('2026年五一假期航班收益分析报告')
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('2026 May Day Holiday Flight Revenue Analysis')
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = MGRAY

spacer(doc)
spacer(doc)

# 标准头部信息
header_items = [
    ('作者', '数字化办公室-AI'),
    ('日期', '2026-05-08'),
    ('数据来源', '航班收益快报明细.xlsx'),
    ('分析模型', 'Qwen3-VL-235B-A22B-Instruct-AWQ（本地私有化大模型）'),
    ('输出目录', 'data-analysis-local/航班收益五一分析-20260508/'),
]
for label, value in header_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}：')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = DKGRAY
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 一,执行摘要（本地模型生成）
# ══════════════════════════════════════════════════════════════
heading1(doc, '一,执行摘要')

exec_summary = """2026年五一假期期间，公司航班运营实现强劲增长：航班架次与旅客量同比分别增长8.7%与11.2%，客座率维持高位（94.5%），营业收入达6.49亿元，同比增长19.3%，显著优于运力增速。然而，受成本结构刚性及部分航线收益承压影响，整体航班收益同比下滑26.6%，盈利表现未达预期。国际航线收入贡献提升，但部分热门国际航线出现亏损，需优化收益管理策略。"""
body(doc, exec_summary)

# KPI总览
heading2(doc, '核心指标总览')
make_table(doc,
    ['指标', '数值', '同比变化', '评估'],
    [['营业收入', fmt(total_rev), f'{rev_chg:+.1f}%', '强劲增长'],
     ['航班架次', f'{total_flights:,}', f'{flight_chg:+.1f}%', '稳步增长'],
     ['旅客人数', f'{total_pax:,}', f'{pax_chg:+.1f}%', '两位数增长'],
     ['平均客座率', f'{avg_load:.1f}%', f'{avg_load-last_load:+.1f}pct', '高位稳定'],
     ['航班收益', fmt(total_profit), f'{profit_chg:+.1f}%', '收益承压']],
    col_widths=[Cm(3), Cm(3), Cm(2.5), Cm(3)])

# ══════════════════════════════════════════════════════════════
# 二,关键发现（本地模型生成）
# ══════════════════════════════════════════════════════════════
heading1(doc, '二,关键发现')

findings = [
    ("收入增长优于运力扩张，但利润承压",
     "收入增速（+19.3%）显著高于航班架次增速（+8.7%），反映票价或附加服务收入提升；但航班收益同比下滑26.6%，主因航油成本占比高达46.1%，叠加起降费与维修成本刚性，压缩利润空间。"),
    ("国际航线贡献突出但收益分化严重",
     "国际航线收入占比24.6%，但'浦东-曼谷'航线出现亏损（-48万），而'大阪-浦东''浦东-大阪'等航线收益优异，显示国际航线需精细化收益管理与成本控制。"),
    ("出行节奏清晰，高峰日盈利能力分化明显",
     "出发与返程高峰日（4.30-5.1,5.5-5.6）客座率超95%，收入贡献超60%，但5月2-3日假期低谷期出现收益负值，反映淡季定价策略或成本结构需优化。"),
    ("核心基地与航线表现稳健",
     "浦东与虹桥基地合计贡献36.1%收入，客座率均超94%，显示枢纽优势；TOP10航线中7条为国内干线或国际热门航线，收益效率较高，但需警惕部分国际航线亏损风险。"),
]

for title, content in findings:
    p = doc.add_paragraph()
    run1 = p.add_run(f'• {title}：')
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = DKGRAY
    run2 = p.add_run(content)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════
# 三,每日运营数据
# ══════════════════════════════════════════════════════════════
heading1(doc, '三,每日运营明细')
make_table(doc,
    ['日期', '架次', '营业收入', '旅客人数', '客座率', '航班收益'],
    [[str(r['航班日期(YYYY-MM-DD)'])[-5:],
      f"{r['架次']:.0f}", fmt(r['收入']),
      f"{r['人数']:,}", f"{r['客座率']:.1f}%", fmt(r['收益'])]
     for _, r in daily.iterrows()],
    col_widths=[Cm(2.2), Cm(2), Cm(3), Cm(2.5), Cm(2), Cm(3)])
spacer(doc)

# 出行节奏
heading2(doc, '出行节奏分析')
make_table(doc,
    ['时段', '日期', '架次', '营业收入', '旅客人数', '平均客座率'],
    [['出发高峰', '4.30-5.1', f"{int(depart['架次'].sum()):.0f}",
      fmt(depart['总收入'].sum()), f"{int(depart['总人数'].sum()):,}",
      f"{depart['客座率'].mean():.1f}%"],
     ['假期低谷', '5.2-5.3', f"{int(mid_df['架次'].sum()):.0f}",
      fmt(mid_df['总收入'].sum()), f"{int(mid_df['总人数'].sum()):,}",
      f"{mid_df['客座率'].mean():.1f}%"],
     ['返程高峰', '5.5-5.6', f"{int(return_df['架次'].sum()):.0f}",
      fmt(return_df['总收入'].sum()), f"{int(return_df['总人数'].sum()):,}",
      f"{return_df['客座率'].mean():.1f}%"]],
    col_widths=[Cm(2.5), Cm(2.5), Cm(2.5), Cm(3), Cm(2.5), Cm(2.5)])

# ══════════════════════════════════════════════════════════════
# 四,航线分析
# ══════════════════════════════════════════════════════════════
heading1(doc, '四,航线与基地分析')

heading2(doc, '4.1 航线性质分类')
make_table(doc,
    ['航线性质', '航班架次', '营业收入', '旅客人数', '平均客座率', '收入占比'],
    [[r['性质'], f"{r['架次']:.0f}", fmt(r['收入']),
      f"{r['人数']:,}", f"{r['客座率']:.1f}%", f"{r['收入']/total_rev*100:.1f}%"]
     for _, r in route_nat.iterrows()],
    col_widths=[Cm(2.5), Cm(2.5), Cm(3), Cm(2.5), Cm(2.5), Cm(2)])
spacer(doc)

heading2(doc, '4.2 基地排名TOP5')
make_table(doc,
    ['排名', '基地', '架次', '营业收入', '旅客人数', '平均客座率'],
    [[f'{i}', r['基地名称'], f"{r['架次']:.0f}",
      fmt(r['收入']), f"{r['人数']:,}", f"{r['客座率']:.1f}%"]
     for i, (_, r) in enumerate(base_top5.iterrows(), 1)],
    col_widths=[Cm(1.2), Cm(3), Cm(2.5), Cm(3), Cm(2.5), Cm(2.5)])
spacer(doc)

heading2(doc, '4.3 航线排名TOP10')
make_table(doc,
    ['排名', '航线', '架次', '营业收入', '旅客人数', '客座率', '航班收益'],
    [[f'{i}', r['航线名称'][:14], f"{r['架次']:.0f}",
      fmt(r['收入']), f"{r['人数']:,}",
      f"{r['客座率']:.1f}%", fmt(r['航班收益'])]
     for i, (_, r) in enumerate(top10.iterrows(), 1)],
    col_widths=[Cm(1), Cm(3), Cm(2), Cm(2.8), Cm(2), Cm(2), Cm(2.5)])

# ══════════════════════════════════════════════════════════════
# 五,成本结构
# ══════════════════════════════════════════════════════════════
heading1(doc, '五,成本结构分析')
make_table(doc,
    ['成本项目', '金额', '占总成本比', '备注'],
    [['变动成本', fmt(var_cost), f'{var_cost/total_cost*100:.1f}%', '燃油,起降,维修等'],
     ['固定成本费用', fmt(fix_cost), f'{fix_cost/total_cost*100:.1f}%', '固定小时成本×轮档小时'],
     ['航油费', fmt(fuel), f'{fuel/total_cost*100:.1f}%', '主要成本驱动因素'],
     ['起降费', fmt(landing), f'{landing/total_cost*100:.1f}%', '含航路费+机场费+旅客服务费'],
     ['维修成本', fmt(maint), f'{maint/total_cost*100:.1f}%', '按轮档小时摊销'],
     ['成本合计', fmt(total_cost), '100.0%', '变动+固定成本合计']],
    col_widths=[Cm(3), Cm(3.5), Cm(2.5), Cm(6)])
spacer(doc)
note(doc, '▶ 成本关注：变动成本占总成本75.8%，航油费占成本总额46.1%，油价波动对收益影响显著，需持续关注国际油价走势。')

# ══════════════════════════════════════════════════════════════
# 六,风险提示（本地模型生成）
# ══════════════════════════════════════════════════════════════
heading1(doc, '六,风险提示')

risks = [
    ("成本刚性风险",
     "航油成本占总成本46.1%，受国际油价波动影响大；固定成本占比24.2%，在低峰期难以摊薄，导致5月2-3日出现收益亏损，需加强成本弹性管理。"),
    ("国际航线收益结构失衡",
     "部分国际航线（如浦东-曼谷）虽客座率高（98.5%），但收益为负，反映票价或成本结构不合理，存在高上座,低收益陷阱，需重新评估航线定价模型。"),
    ("返程高峰后收益回落",
     "5月6日虽客座率高（95.4%），但收益仅510万，较5月5日（2,968万）大幅下滑，显示返程后期定价或需求管理存在优化空间。"),
]

for title, content in risks:
    p = doc.add_paragraph()
    run1 = p.add_run(f'• {title}：')
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = RED
    run2 = p.add_run(content)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════
# 七,下一步建议（本地模型生成）
# ══════════════════════════════════════════════════════════════
heading1(doc, '七,下一步建议')

recs = [
    ("优化国际航线收益管理",
     "对亏损国际航线（如浦东-曼谷）进行成本-收益重构，考虑动态定价,舱位控制或与航司联营，提升单航线盈利能力；同时复制大阪-浦东等高收益航线的成功模式。"),
    ("强化低峰期成本控制与定价策略",
     "针对5月2-3日等低谷期，可推行早鸟折扣+动态调价机制，同时通过包机或促销提升利用率，降低单位成本摊薄压力。"),
    ("推进航油成本对冲机制",
     "建议财务部联合采购部门，探索航油期货或期权对冲工具，锁定部分成本，缓解油价波动对利润的冲击。"),
    ("深化基地与航线绩效评估",
     "建立收入-收益-成本三位一体的航线绩效仪表盘，按周/月动态监控TOP航线与基地表现，支持资源向高收益航线倾斜，提升整体资产回报率。"),
]

for i, (title, content) in enumerate(recs, 1):
    p = doc.add_paragraph()
    run1 = p.add_run(f'{i}. {title}：')
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = BLUE
    run2 = p.add_run(content)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

spacer(doc)

# 结语
conclusion = "五一假期整体运营表现优异，收入增长强劲，但利润端受成本制约。建议公司聚焦收益精细化管理与成本弹性优化，以实现高上座,高收益的可持续增长目标。"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run(conclusion)
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = DKGRAY

spacer(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run('本报告由本地大模型分析系统自动生成  |  数据来源：航班收益快报明细.xlsx  |  分析日期：2026-05-08')
fr.font.size  = Pt(8)
fr.font.italic = True
fr.font.color.rgb = MGRAY

# ─── 保存 ───────────────────────────────────────────────
out = '/Users/fox/Claude Code/data-analysis-local/航班收益五一分析-20260508/航班收益快报_五一分析报告_本地模型版.docx'
doc.save(out)
print(f'Word报告已生成: {out}')